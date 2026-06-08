# -*- coding: utf-8 -*-
"""md2word — 通用 House 风格 Markdown→Word 转换器（CLI）。
特性：内置 Heading 样式（标题导航窗格）/ 专业表格（深蓝表头+斑马纹+紧凑单倍行距单元格）/
代码块阴影 / 引用 callout / 列表 / <br> 换行 / 可选把指定 fenced block 替换为图片。
用法：
  python md2word.py input.md [-o output.docx] [--title T] [--subtitle S] [--org 公司名]
                    [--img-dir DIR --map map.json]   # map.json: {"块内出现的签名": "图片名(不含扩展)"}
依赖：pip install python-docx --break-system-packages
"""
import sys, re, os, argparse, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="1F4E79"; NAVY2="2E5B8A"; SUB="D6E4F0"; ZEBRA="F2F6FB"
CODEBG="F4F4F4"; CALLBG="EAF1FB"
CJK="Microsoft YaHei"; MONO="Consolas"; CODEFONT="NSimSun"

def set_font(run, name):
    run.font.name=name
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),name)

def tighten(p):
    pf=p.paragraph_format
    pf.space_before=Pt(0); pf.space_after=Pt(0)
    pf.line_spacing=1.0; pf.line_spacing_rule=WD_LINE_SPACING.SINGLE

def shade(el, hexc):
    pr=el.get_or_add_tcPr() if el.tag==qn('w:tc') else el.get_or_add_pPr()
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc)
    pr.append(sh)

def cell_borders(cell, color="BBBBBB", sz=4):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color); b.append(e)
    tcPr.append(b)

def para_left_bar(p, color):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr')
    l=OxmlElement('w:left'); l.set(qn('w:val'),'single'); l.set(qn('w:sz'),'18'); l.set(qn('w:space'),'8'); l.set(qn('w:color'),color)
    pbdr.append(l); pPr.append(pbdr)

def heading_bottom(p, color):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),color)
    pbdr.append(b); pPr.append(pbdr)

INLINE=re.compile(r'(\*\*.+?\*\*|`.+?`|\[\[.+?\]\])')
def add_runs(p, text, size=10.5, color=None, base_bold=False):
    text=re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    for pidx,part in enumerate(re.split(r'<br\s*/?>', text)):
        if pidx>0:
            br=p.add_run(); br.add_break()
        for seg in INLINE.split(part):
            if not seg: continue
            b=base_bold; mono=False; t=seg
            if seg.startswith('**') and seg.endswith('**'): b=True; t=seg[2:-2]
            elif seg.startswith('`') and seg.endswith('`'): mono=True; t=seg[1:-1]
            elif seg.startswith('[[') and seg.endswith(']]'): t=seg[2:-2]
            r=p.add_run(t); r.bold=b; r.font.size=Pt(size)
            set_font(r, MONO if mono else CJK)
            if color: r.font.color.rgb=RGBColor.from_string(color)
    return p

def img_dims(png):
    import struct
    d=open(png,'rb').read(33); w,h=struct.unpack('>II',d[16:24]); return w,h

class Conv:
    def __init__(self, doc, diagram_map, imgdir):
        self.doc=doc; self.dmap=diagram_map or {}; self.imgdir=imgdir
    def add_image(self, key):
        png=os.path.join(self.imgdir,key+'.png'); svg=os.path.join(self.imgdir,key+'.svg')
        w,h=img_dims(png); wi=min(6.4, w/150.0)
        p=self.doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run()
        try: run.add_picture(svg if os.path.exists(svg) else png, width=Inches(wi))
        except Exception: run.add_picture(png, width=Inches(wi))
    def heading(self, level, text):
        try: p=self.doc.add_paragraph(style='Heading %d'%min(level,5))
        except KeyError: p=self.doc.add_paragraph()
        p.paragraph_format.space_before=Pt(13 if level<=2 else 9); p.paragraph_format.space_after=Pt(6)
        size={1:16,2:14,3:12.5,4:11.5}.get(level,11); col=NAVY if level<=2 else NAVY2
        add_runs(p, text, size=size, color=col, base_bold=True)
        if level==1: heading_bottom(p, NAVY)
    def para(self, text):
        p=self.doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); add_runs(p, text, size=10.5)
    def bullet(self, text, lvl=0):
        p=self.doc.add_paragraph(style='List Bullet' if lvl==0 else 'List Bullet 2')
        p.paragraph_format.space_after=Pt(3); add_runs(p, text, size=10.5)
    def numbered(self, text):
        p=self.doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); add_runs(p, text, size=10.5)
    def callout(self, text):
        p=self.doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
        shade(p._p, CALLBG); para_left_bar(p, NAVY); p.paragraph_format.left_indent=Pt(8); add_runs(p, text, size=10)
    def code_block(self, lines):
        for i,l in enumerate(lines):
            p=self.doc.add_paragraph(); shade(p._p, CODEBG); para_left_bar(p,"9BB4CC")
            p.paragraph_format.left_indent=Pt(8)
            p.paragraph_format.space_before=Pt(2 if i==0 else 0); p.paragraph_format.space_after=Pt(2 if i==len(lines)-1 else 0)
            r=p.add_run(l if l.strip() else ' '); r.font.size=Pt(8.5); set_font(r, CODEFONT)
    def table(self, rows):
        header=rows[0]; body=rows[1:]
        t=self.doc.add_table(rows=0, cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
        hr=t.add_row().cells
        for i,htxt in enumerate(header):
            cell_borders(hr[i]); shade(hr[i]._tc, NAVY)
            pr=hr[i].paragraphs[0]; pr.alignment=WD_ALIGN_PARAGRAPH.CENTER; tighten(pr)
            add_runs(pr, htxt, size=9.5, color="FFFFFF", base_bold=True)
        for ri,row in enumerate(body):
            cells=t.add_row().cells
            for i in range(len(header)):
                txt=row[i] if i<len(row) else ""
                cell_borders(cells[i])
                if ri%2==1: shade(cells[i]._tc, ZEBRA)
                pr=cells[i].paragraphs[0]; tighten(pr); add_runs(pr, txt, size=9)
        self.doc.add_paragraph().paragraph_format.space_after=Pt(2)

def parse(md_path, doc, diagram_map, imgdir):
    conv=Conv(doc, diagram_map, imgdir)
    raw=open(md_path,encoding='utf-8').read()
    raw=re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]','',raw)
    lines=raw.split('\n'); i=0; n=len(lines)
    if lines and lines[0].strip()=='---':
        j=1
        while j<n and lines[j].strip()!='---': j+=1
        i=j+1
    while i<n:
        line=lines[i]; s=line.strip()
        if not s: i+=1; continue
        if s.startswith('```'):
            j=i+1; buf=[]
            while j<n and not lines[j].strip().startswith('```'): buf.append(lines[j]); j+=1
            content='\n'.join(buf); key=None
            for sig,k in (diagram_map or {}).items():
                if sig in content: key=k; break
            if key: conv.add_image(key)
            else: conv.code_block(buf)
            i=j+1; continue
        if s.startswith('|') and i+1<n and re.match(r'^\s*\|?[\s:\-\|]+\|?\s*$', lines[i+1]):
            trows=[]; k=i
            while k<n and lines[k].strip().startswith('|'): trows.append(lines[k]); k+=1
            parsed=[]
            for idx,rl in enumerate(trows):
                if idx==1: continue
                parsed.append([c.strip() for c in rl.strip().strip('|').split('|')])
            conv.table(parsed); i=k; continue
        m=re.match(r'^(#{1,6})\s+(.*)$', s)
        if m: conv.heading(len(m.group(1)), m.group(2).strip()); i+=1; continue
        if re.match(r'^---+$', s): i+=1; continue
        if s.startswith('>'):
            buf=[]
            while i<n and lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?','',lines[i])); i+=1
            conv.callout('<br>'.join(x for x in buf if x.strip())); continue
        mb=re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if mb: conv.bullet(mb.group(2), 1 if len(mb.group(1))>=2 else 0); i+=1; continue
        mn=re.match(r'^\s*\d+\.\s+(.*)$', line)
        if mn: conv.numbered(mn.group(1)); i+=1; continue
        conv.para(s); i+=1

def setup_heading_styles(doc):
    specs={1:(16,NAVY),2:(14,NAVY2),3:(12.5,NAVY2),4:(11.5,NAVY2),5:(11,NAVY2)}
    for lvl,(sz,col) in specs.items():
        try: stl=doc.styles['Heading %d'%lvl]
        except KeyError: continue
        stl.font.name=CJK; stl.font.size=Pt(sz); stl.font.bold=True
        stl.font.color.rgb=RGBColor.from_string(col)
        rpr=stl.element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
        if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),CJK)

def set_font_default(doc):
    rpr=doc.styles['Normal'].element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),CJK)

def add_footer(sec):
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    def fld(p, instr):
        r=p.add_run(); fb=OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'),'begin'); r._r.append(fb)
        r2=p.add_run(); it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr; r2._r.append(it)
        r3=p.add_run(); fe=OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'),'end'); r3._r.append(fe)
    r=fp.add_run("第 "); set_font(r,CJK); fld(fp,"PAGE")
    r=fp.add_run(" 页 / 共 "); set_font(r,CJK); fld(fp,"NUMPAGES")
    r=fp.add_run(" 页"); set_font(r,CJK)
    for rr in fp.runs: rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string("999999")

def derive_title(md_path):
    try:
        for l in open(md_path,encoding='utf-8'):
            m=re.match(r'^#{1,2}\s+(.*)$', l.strip())
            if m: return m.group(1).strip()
    except Exception: pass
    return os.path.splitext(os.path.basename(md_path))[0]

PPR_ORDER=['pStyle','keepNext','keepLines','pageBreakBefore','framePr','widowControl','numPr',
 'suppressLineNumbers','pBdr','shd','tabs','suppressAutoHyphens','kinsoku','wordWrap','overflowPunct',
 'topLinePunct','autoSpaceDE','autoSpaceDN','bidi','adjustRightInd','snapToGrid','spacing','ind',
 'contextualSpacing','mirrorIndents','suppressOverlap','jc','textDirection','textAlignment',
 'textboxTightWrap','outlineLvl','divId','cnfStyle','rPr','sectPr','pPrChange']
def reorder_ppr(doc):
    idx={n:i for i,n in enumerate(PPR_ORDER)}
    for pPr in doc.element.iter(qn('w:pPr')):
        ch=list(pPr); ch.sort(key=lambda e: idx.get(e.tag.split('}')[-1],999))
        for c in ch: pPr.remove(c)
        for c in ch: pPr.append(c)

def build(md_path, out_path, title=None, subtitle="", org="", diagram_map=None, imgdir="."):
    title=title or derive_title(md_path)
    doc=Document()
    z=doc.settings.element.find(qn('w:zoom'))
    if z is not None and z.get(qn('w:percent')) is None: z.set(qn('w:percent'),'100')
    st=doc.styles['Normal']; st.font.name=CJK; st.font.size=Pt(10.5)
    set_font_default(doc); setup_heading_styles(doc)
    sec=doc.sections[0]
    for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,m,Inches(1))
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=hp.add_run(title); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string("999999"); set_font(r,CJK)
    add_footer(sec)
    if org:
        tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=tp.add_run(org); r.bold=True; r.font.size=Pt(12); set_font(r,CJK)
    tp2=doc.add_paragraph(); tp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp2.add_run(title); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(NAVY); set_font(r,CJK)
    if subtitle:
        sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=sp.add_run(subtitle); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string("555555"); set_font(r,CJK)
    doc.add_paragraph()
    parse(md_path, doc, diagram_map, imgdir)
    reorder_ppr(doc)
    doc.save(out_path); print("saved", out_path)

def main():
    ap=argparse.ArgumentParser(description="House 风格 Markdown→Word 转换器")
    ap.add_argument("input"); ap.add_argument("-o","--output", default=None)
    ap.add_argument("--title", default=None); ap.add_argument("--subtitle", default="")
    ap.add_argument("--org", default="")
    ap.add_argument("--img-dir", default="."); ap.add_argument("--map", default=None)
    a=ap.parse_args()
    out=a.output or os.path.splitext(a.input)[0]+".docx"
    dmap=json.load(open(a.map,encoding='utf-8')) if a.map else None
    build(a.input, out, a.title, a.subtitle, a.org, dmap, a.img_dir)

if __name__=="__main__": main()
