import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

def tighten_table(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf=p.paragraph_format
                pf.space_before=Pt(0); pf.space_after=Pt(0)
                pf.line_spacing=1.0; pf.line_spacing_rule=WD_LINE_SPACING.SINGLE
            for nt in cell.tables:   # 嵌套表
                tighten_table(nt)

d=Document(sys.argv[1])
for t in d.tables: tighten_table(t)
d.save(sys.argv[2])
print("tightened ->", sys.argv[2])
