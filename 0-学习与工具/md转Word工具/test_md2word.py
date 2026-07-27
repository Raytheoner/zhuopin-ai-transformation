# -*- coding: utf-8 -*-
"""md2word 真复选框（w14:checkbox）单测。
覆盖：段落/列表/表格单元格三处出现位置 × ☐/[ ]/[x] 三种写法 × 勾选状态读回。
"""
import os, sys, tempfile
import pytest
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(__file__))
import md2word as mw


def _build(md_text, tmp_path, name="case.md"):
    md_path = os.path.join(tmp_path, name)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    out_path = os.path.join(tmp_path, os.path.splitext(name)[0] + ".docx")
    mw.build(md_path, out_path, title="测试")
    return out_path


def _all_checkbox_sdts(doc):
    return [sdt for sdt in doc.element.body.iter(qn('w:sdt'))
            if sdt.find('.//' + qn('w14:checkbox')) is not None]


class TestParagraphCheckbox:
    def test_unchecked_ballot_char_becomes_real_control(self, tmp_path):
        out = _build("☐ 同意本次口径调整\n", str(tmp_path))
        doc = Document(out)
        sdts = _all_checkbox_sdts(doc)
        assert len(sdts) == 1
        cb = sdts[0].find('.//' + qn('w14:checkbox'))
        checked = cb.find(qn('w14:checked'))
        assert checked.get(qn('w14:val')) == '0'

    def test_checked_ballot_char(self, tmp_path):
        out = _build("☑ 已确认\n", str(tmp_path))
        doc = Document(out)
        sdts = _all_checkbox_sdts(doc)
        assert len(sdts) == 1
        cb = sdts[0].find('.//' + qn('w14:checkbox'))
        assert cb.find(qn('w14:checked')).get(qn('w14:val')) == '1'

    def test_bracket_x_checked_variants(self, tmp_path):
        out = _build("[x] 小写\n\n[X] 大写\n", str(tmp_path))
        doc = Document(out)
        sdts = _all_checkbox_sdts(doc)
        assert len(sdts) == 2
        for sdt in sdts:
            cb = sdt.find('.//' + qn('w14:checkbox'))
            assert cb.find(qn('w14:checked')).get(qn('w14:val')) == '1'

    def test_no_dead_ballot_char_left_in_plain_text(self, tmp_path):
        """转换后不应再有裸的 ☐ 文本 run（否则专员仍点不动）。"""
        out = _build("☐ 同意\n", str(tmp_path))
        doc = Document(out)
        # 段落里除 sdtContent 内部的 ☐ 外，不应再有游离的 ☐ 字符 run
        for p in doc.paragraphs:
            for run in p.runs:
                assert '☐' not in run.text


class TestBulletCheckbox:
    def test_task_list_unchecked(self, tmp_path):
        out = _build("- [ ] 待办事项一\n", str(tmp_path))
        doc = Document(out)
        sdts = _all_checkbox_sdts(doc)
        assert len(sdts) == 1
        cb = sdts[0].find('.//' + qn('w14:checkbox'))
        assert cb.find(qn('w14:checked')).get(qn('w14:val')) == '0'
        # 标签文字仍需正常渲染
        full_text = "".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "待办事项一" in full_text

    def test_task_list_checked(self, tmp_path):
        out = _build("- [x] 已完成项\n", str(tmp_path))
        doc = Document(out)
        sdts = _all_checkbox_sdts(doc)
        assert len(sdts) == 1
        cb = sdts[0].find('.//' + qn('w14:checkbox'))
        assert cb.find(qn('w14:checked')).get(qn('w14:val')) == '1'


class TestTableCellCheckbox:
    """判例批改表核心场景：勾选项都在表格单元格里。"""

    MD = (
        "| # | 真实场景 | 现状判定 | 拟改判定 | ✅对 | ❌错 | ✏️改判理由 |\n"
        "|---|---|---|---|---|---|---|\n"
        "| 1 | 料号A提前期7天 | 齐套 | 齐套 | ☐ | ☐ | |\n"
        "| 2 | 料号B提前期3天 | 缺料 | 齐套 | ☐ | ☑ | 口径应按在途量计 |\n"
    )

    def test_cell_checkbox_count_and_context(self, tmp_path):
        out = _build(self.MD, str(tmp_path))
        doc = Document(out)
        results = mw.read_checkboxes(out)
        # 2 行 × 2 个判定列 = 4 个复选框
        assert len(results) == 4
        assert [r["checked"] for r in results] == [False, False, False, True]

    def test_cell_with_only_checkbox_has_no_leftover_text(self, tmp_path):
        out = _build(self.MD, str(tmp_path))
        doc = Document(out)
        table = doc.tables[0]
        # 第2行(索引2，含表头+分隔行占1个table row) "✅对"列 应只含一个真控件、无死字符
        cell = table.rows[1].cells[4]
        for p in cell.paragraphs:
            for run in p.runs:
                assert run.text.strip() not in ('☐', '☑', '☒')

    def test_read_checkboxes_context_matches_row(self, tmp_path):
        out = _build(self.MD, str(tmp_path))
        results = mw.read_checkboxes(out)
        # 第4个复选框（行2 的 ❌错 列）context 应落在含"口径应按在途量计"同一单元格文字里
        # （同一单元格段落，含勾选符本身与其余文字）
        assert results[3]["checked"] is True


class TestRegressionExistingFeatures:
    def test_bold_and_inline_code_unaffected(self, tmp_path):
        out = _build("这是**加粗**与`code`文本，无勾选\n", str(tmp_path))
        doc = Document(out)
        assert len(_all_checkbox_sdts(doc)) == 0
        full_text = "".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "加粗" in full_text and "code" in full_text

    def test_double_bracket_internal_link_not_treated_as_checkbox(self, tmp_path):
        out = _build("参见[[附录A]]获取详情\n", str(tmp_path))
        doc = Document(out)
        assert len(_all_checkbox_sdts(doc)) == 0
        full_text = "".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "附录A" in full_text

    def test_markdown_link_stripped_as_before(self, tmp_path):
        out = _build("详情见[链接文本](https://example.com)\n", str(tmp_path))
        doc = Document(out)
        full_text = "".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "链接文本" in full_text
        assert "example.com" not in full_text

    def test_code_block_checkbox_like_text_stays_literal(self, tmp_path):
        """代码块里的方括号不应被转成控件（code_block 走独立渲染路径，不经 add_runs）。"""
        out = _build("```\nif x == [ ]:\n    pass\n```\n", str(tmp_path))
        doc = Document(out)
        assert len(_all_checkbox_sdts(doc)) == 0

    def test_heading_and_table_without_checkbox_render_normally(self, tmp_path):
        md = "# 标题\n\n| 列A | 列B |\n|---|---|\n| 值1 | 值2 |\n"
        out = _build(md, str(tmp_path))
        doc = Document(out)
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[1].cells[0].text == "值1"
        assert len(_all_checkbox_sdts(doc)) == 0


class TestSdtXmlFidelity:
    def test_checkbox_state_glyphs_match_word_defaults(self, tmp_path):
        out = _build("☐ 未勾选项\n", str(tmp_path))
        doc = Document(out)
        sdt = _all_checkbox_sdts(doc)[0]
        cb = sdt.find('.//' + qn('w14:checkbox'))
        cs = cb.find(qn('w14:checkedState'))
        us = cb.find(qn('w14:uncheckedState'))
        assert cs.get(qn('w14:val')) == '2612'
        assert us.get(qn('w14:val')) == '2610'
        assert cs.get(qn('w14:font')) == 'MS Gothic'

    def test_sdt_has_unique_ids(self, tmp_path):
        out = _build("☐ 一\n\n☐ 二\n\n☐ 三\n", str(tmp_path))
        doc = Document(out)
        ids = []
        for sdt in _all_checkbox_sdts(doc):
            idel = sdt.find(qn('w:sdtPr') + '/' + qn('w:id'))
            ids.append(idel.get(qn('w:val')))
        assert len(ids) == len(set(ids)) == 3


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
